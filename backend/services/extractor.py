"""
backend/services/extractor.py — Text extraction services for PDF, DOCX, and TXT files
"""

import io
import logging
from typing import Tuple, Dict, Any
import PyPDF2
import docx

logger = logging.getLogger(__name__)


class TextExtractionError(Exception):
    """Custom exception raised when text extraction fails."""
    pass


def extract_text_from_pdf(file_stream) -> Tuple[str, int]:
    """
    Extract text and page count from a PDF file stream.
    """
    try:
        file_stream.seek(0)
        # Wrap in BytesIO in case stream is not fully compatible with PyPDF2
        pdf_bytes = io.BytesIO(file_stream.read())
        reader = PyPDF2.PdfReader(pdf_bytes)
        page_count = len(reader.pages)
        
        extracted_text_list = []
        for i in range(page_count):
            page = reader.pages[i]
            page_text = page.extract_text()
            if page_text:
                extracted_text_list.append(page_text)
                
        full_text = "\n\n".join(extracted_text_list)
        return full_text, page_count
    except Exception as e:
        logger.error("Failed to extract text from PDF: %s", str(e), exc_info=True)
        raise TextExtractionError(f"Failed to parse PDF file: {str(e)}")


def extract_text_from_docx(file_stream) -> Tuple[str, int]:
    """
    Extract text and estimate page count from a DOCX file stream.
    """
    try:
        file_stream.seek(0)
        docx_bytes = io.BytesIO(file_stream.read())
        doc = docx.Document(docx_bytes)
        
        paragraphs_text = [p.text for p in doc.paragraphs if p.text]
        
        # Also extract text from tables
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                row_cells = [cell.text for cell in row.cells if cell.text]
                if row_cells:
                    tables_text.append(" | ".join(row_cells))
                    
        full_text = "\n".join(paragraphs_text + tables_text)
        
        # Estimate page count (1 page approx. 500 words, minimum 1)
        words = len(full_text.split())
        estimated_pages = max(1, (words + 499) // 500)
        
        return full_text, estimated_pages
    except Exception as e:
        logger.error("Failed to extract text from DOCX: %s", str(e), exc_info=True)
        raise TextExtractionError(f"Failed to parse Word Document: {str(e)}")


def extract_text_from_txt(file_stream) -> Tuple[str, int]:
    """
    Extract text from a TXT file stream with encoding fallback.
    """
    try:
        file_stream.seek(0)
        content_bytes = file_stream.read()
        
        # Try decoding with UTF-8, fallback to Latin-1
        try:
            full_text = content_bytes.decode('utf-8')
        except UnicodeDecodeError:
            logger.warning("UTF-8 decoding failed, falling back to Latin-1")
            full_text = content_bytes.decode('latin-1')
            
        return full_text, 1
    except Exception as e:
        logger.error("Failed to extract text from TXT: %s", str(e), exc_info=True)
        raise TextExtractionError(f"Failed to parse text file: {str(e)}")


def extract_metadata_and_text(file_stream, filename: str) -> Dict[str, Any]:
    """
    Orchestrate text extraction based on file extension and return text + counts.
    """
    ext = filename.rsplit('.', 1)[1].lower() if '.' in filename else ''
    
    if ext == 'pdf':
        text, page_count = extract_text_from_pdf(file_stream)
    elif ext == 'docx':
        text, page_count = extract_text_from_docx(file_stream)
    elif ext == 'txt':
        text, page_count = extract_text_from_txt(file_stream)
    else:
        raise TextExtractionError(f"Extraction not supported for extension: {ext}")
        
    word_count = len(text.split())
    character_count = len(text)
    
    return {
        "text": text,
        "page_count": page_count,
        "word_count": word_count,
        "character_count": character_count
    }
